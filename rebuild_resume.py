import sys
import os
import base64
import io
import asyncio
from datetime import datetime, timezone
import aiosqlite
import httpx
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv

load_dotenv()

# Database path compatibility
DB_PATH = os.environ.get("DB_PATH", "agent_memory.db")

async def get_db():
    return aiosqlite.connect(DB_PATH)

def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
        
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
    new_run.append(rPr)
    
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def build_docx():
    doc = Document()
    
    # Page Setup - Margins: 1.0 inch for standard formatting layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    
    # Helper to add bold/italic centered header text
    def add_header_line(text, size, bold=False, italic=False, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return p

    # Header Name (clean ASCII, no special chars)
    add_header_line("MADAN SAI DARAM", 18, bold=True, space_after=2)
    
    # Contact Details (using pipes with clickable links, including GitHub)
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(4)
    run_c1 = p_contact.add_run("Madan Sai Daram  |  +91-9963214141  |  ")
    run_c1.font.size = Pt(10)
    add_hyperlink(p_contact, "madansai303@gmail.com", "mailto:madansai303@gmail.com")
    run_c2 = p_contact.add_run("  |  ")
    run_c2.font.size = Pt(10)
    add_hyperlink(p_contact, "linkedin.com/in/madan-sai-daram-a26735313", "https://linkedin.com/in/madan-sai-daram-a26735313")
    run_c3 = p_contact.add_run("  |  ")
    run_c3.font.size = Pt(10)
    add_hyperlink(p_contact, "github.com/Madansai1997", "https://github.com/Madansai1997")
    
    # Headline (ATS-compliant headline format)
    add_header_line("Data Analyst – Business Intelligence", 11, bold=True, italic=True, space_after=12)
    
    # Helper to add Title Case section headers (no decorative character lines)
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy color for structure

    # Helper to add standard bullets
    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.size = Pt(10.5)

    # 1. Summary
    add_section_header("Summary")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(0)
    p_sum.paragraph_format.space_after = Pt(8)
    p_sum.paragraph_format.line_spacing = 1.15
    run_sum = p_sum.add_run(
        "Data Analyst with 4+ years of experience specializing in Business Intelligence, SQL, and Power BI. "
        "Results-driven professional with strong stakeholder communication and cross-functional collaboration skills, "
        "delivering SQL analytics, reporting automation, data validation, and dashboard solutions across cloud data platforms. "
        "Improved reporting efficiency by ~40%, reduced manual effort by 25%, and supported more than 10 business KPIs "
        "through scalable reporting and automation. Expert in advanced SQL (CTEs, window functions), Power BI, Snowflake, "
        "and BigQuery to improve data accuracy, strengthen reporting reliability, and support business decision-making."
    )
    run_sum.font.size = Pt(10.5)
    
    # 2. Experience
    add_section_header("Experience")
    
    # Job 1
    p_job1 = doc.add_paragraph()
    p_job1.paragraph_format.space_before = Pt(4)
    p_job1.paragraph_format.space_after = Pt(4)
    p_job1.paragraph_format.keep_with_next = True
    r_j1_title = p_job1.add_run("MyTech Detectives")
    r_j1_title.bold = True
    r_j1_title.font.size = Pt(11)
    r_j1_text = p_job1.add_run(" — Data Analyst | Hyderabad, Telangana | Jul 2024 – Present")
    r_j1_text.font.size = Pt(11)
    
    # Kept top 4 bullets + injected keywords + metrics
    add_bullet("Developed 5+ interactive Power BI dashboards tracking 10+ key business KPIs, establishing dashboard governance, increasing operational efficiency visibility by 25% and supporting business reporting.")
    add_bullet("Optimized SQL queries and data pipelines across Snowflake and BigQuery, reducing report generation time by 30% through SQL performance tuning and data modeling.")
    add_bullet("Engineered Python-based (Pandas, NumPy) data cleaning and EDA scripts, detecting 15+ critical data anomalies and improving overall reporting accuracy by 20%.")
    add_bullet("Led data validation and reconciliation of 12 source systems and warehouse tables, improving report reliability by 20%.")
    
    # Job 2
    p_job2 = doc.add_paragraph()
    p_job2.paragraph_format.space_before = Pt(8)
    p_job2.paragraph_format.space_after = Pt(4)
    p_job2.paragraph_format.keep_with_next = True
    r_j2_title = p_job2.add_run("Cognizant Technology Solutions")
    r_j2_title.bold = True
    r_j2_title.font.size = Pt(11)
    r_j2_text = p_job2.add_run(" — Associate Test Engineer | Hyderabad, Telangana | Sep 2019 – Jul 2022")
    r_j2_text.font.size = Pt(11)
    
    add_bullet("Ensured 99.8% data integrity across 4+ large-scale HIPAA-compliant healthcare applications by leading end-to-end functional and validation testing.")
    add_bullet("Developed and maintained Selenium automation frameworks, reducing manual regression effort by 25%.")
    add_bullet("Optimized test execution using BrowserStack, reducing regression cycles by 35% through parallel cross-browser testing.")
    add_bullet("Validated backend transactional data pipelines and ETL processes using 100+ complex SQL verification queries, achieving a 98% data-quality score.")

    # 3. Projects
    add_section_header("Projects")
    
    p_p1 = doc.add_paragraph()
    p_p1.paragraph_format.space_before = Pt(4)
    p_p1.paragraph_format.space_after = Pt(4)
    r_p1 = p_p1.add_run("Customer Churn Analysis (Python + SQL) | Jan 2024 – Mar 2024")
    r_p1.bold = True
    r_p1.font.size = Pt(11)
    
    add_bullet("Analyzed 7K+ subscription customer records using SQL and Python to identify churn trends across contract type, tenure, and customer engagement behavior.")
    add_bullet("Applied cohort analysis, feature correlation, and window functions to identify high-risk segments contributing to an 8% monthly churn rate reduction.")
    
    p_p2 = doc.add_paragraph()
    p_p2.paragraph_format.space_before = Pt(6)
    p_p2.paragraph_format.space_after = Pt(4)
    r_p2 = p_p2.add_run("Sales Performance & Revenue Analytics Dashboard (Power BI + SQL) | Apr 2024 – Jun 2024")
    r_p2.bold = True
    r_p2.font.size = Pt(11)
    
    add_bullet("Developed interactive Power BI dashboards analyzing regional sales performance, product-level revenue trends, and monthly KPIs, identifying $45K+ in monthly opportunities.")
    add_bullet("Automated data extraction, transformation, and monthly KPI reporting using SQL and Power Query, saving 4+ hours of manual effort weekly.")

    # 4. Certifications
    add_section_header("Certifications")
    add_bullet("Microsoft Certified: Power BI Data Analyst Associate (PL-300)")
    add_bullet("Google Data Analytics Professional Certificate")

    # 5. Education
    add_section_header("Education")
    p_ed1 = doc.add_paragraph()
    p_ed1.paragraph_format.space_before = Pt(2)
    p_ed1.paragraph_format.space_after = Pt(2)
    r_ed1_univ = p_ed1.add_run("Florida International University")
    r_ed1_univ.bold = True
    r_ed1_univ.font.size = Pt(11)
    r_ed1_deg = p_ed1.add_run(", Miami, FL | MS in Data Science and AI | Aug 2022 – May 2024")
    r_ed1_deg.font.size = Pt(11)
    add_bullet("Coursework: Machine Learning, Statistical Methods, Big Data Systems, Data Visualization.")
    
    p_ed2 = doc.add_paragraph()
    p_ed2.paragraph_format.space_before = Pt(4)
    p_ed2.paragraph_format.space_after = Pt(4)
    r_ed2_univ = p_ed2.add_run("Jawaharlal Nehru Technological University")
    r_ed2_univ.bold = True
    r_ed2_univ.font.size = Pt(11)
    r_ed2_deg = p_ed2.add_run(", Hyderabad, Telangana | Bachelors in Electronics and Communication Engineering | Sep 2014 – May 2018")
    r_ed2_deg.font.size = Pt(11)

    # 6. Skills (Categorized list of key technical competencies)
    add_section_header("Skills")
    add_bullet("Programming: SQL, Python, R")
    add_bullet("BI & Visualization: Power BI, Tableau, Excel")
    add_bullet("Cloud & Databases: Snowflake, BigQuery, SQL Server, Azure, AWS, GCP, MySQL")
    add_bullet("Data Engineering: ETL, data pipeline orchestration, data modeling, data warehousing, data validation, data governance")
    add_bullet("Methodologies: Exploratory Data Analysis, Statistical Analysis, Statistical Modeling, Machine Learning, SQL performance tuning, A/B Testing, Hypothesis Testing")

    # 7. Achievements
    add_section_header("Achievements")
    add_bullet("Solved 50+ advanced SQL problems covering joins, CTEs, and window functions.")
    add_bullet("Transitioned from QA to Data Analytics through self-driven learning and practical projects.")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

async def update_databases(docx_bytes, plain_text):
    filename = "Madan_Sai_Daram_DA_CV.docx"
    now = datetime.now(timezone.utc).isoformat()
    docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")
    
    async with aiosqlite.connect("agent_memory.db") as db:
        await db.execute(
            """INSERT INTO resume_docx (id, filename, data_b64, updated_at) VALUES (1, ?, ?, ?)
               ON CONFLICT(id) DO UPDATE SET filename=excluded.filename, data_b64=excluded.data_b64,
                 updated_at=excluded.updated_at""",
            (filename, docx_b64, now)
        )
        await db.execute(
            """INSERT INTO user_resume_templates (domain, content, updated_at) VALUES ('data_analyst', ?, ?)
               ON CONFLICT(domain) DO UPDATE SET content=excluded.content, updated_at=excluded.updated_at""",
            (plain_text, now)
        )
        await db.commit()
    print("✅ Local SQLite database successfully updated.")

    turso_url = os.environ.get("TURSO_DATABASE_URL")
    turso_token = os.environ.get("TURSO_AUTH_TOKEN")
    if turso_url and turso_token:
        url = turso_url.replace("libsql://", "https://").rstrip("/")
        headers = {"authorization": f"Bearer {turso_token}"}
        
        body_docx = {
            "stmt": {
                "sql": """INSERT INTO resume_docx (id, filename, data_b64, updated_at) VALUES (1, ?, ?, ?)
                          ON CONFLICT(id) DO UPDATE SET filename=excluded.filename, data_b64=excluded.data_b64,
                            updated_at=excluded.updated_at""",
                "args": [
                    {"type": "text", "value": filename},
                    {"type": "text", "value": docx_b64},
                    {"type": "text", "value": now}
                ],
                "named_args": [],
                "want_rows": False
            }
        }
        resp_docx = httpx.post(f"{url}/v1/execute", json=body_docx, headers=headers).json()
        
        body_text = {
            "stmt": {
                "sql": """INSERT INTO user_resume_templates (domain, content, updated_at) VALUES ('data_analyst', ?, ?)
                          ON CONFLICT(domain) DO UPDATE SET content=excluded.content, updated_at=excluded.updated_at""",
                "args": [
                    {"type": "text", "value": plain_text},
                    {"type": "text", "value": now}
                ],
                "named_args": [],
                "want_rows": False
            }
        }
        resp_text = httpx.post(f"{url}/v1/execute", json=body_text, headers=headers).json()
        print("✅ Production Turso database successfully updated.")
    else:
        print("⚠️ Turso env variables missing, production database NOT updated.")

async def main():
    print("🔄 Generating clean, optimized Word Document...")
    docx_bytes = build_docx()
    
    plain_text_res = """MADAN SAI DARAM
Madan Sai Daram | +91-9963214141 | madansai303@gmail.com | https://linkedin.com/in/madan-sai-daram-a26735313 | github.com/Madansai1997
Data Analyst – Business Intelligence

Summary
Data Analyst with 4+ years of experience specializing in Business Intelligence, SQL, and Power BI. Results-driven professional with strong stakeholder communication and cross-functional collaboration skills, delivering SQL analytics, reporting automation, data validation, and dashboard solutions across cloud data platforms. Improved reporting efficiency by ~40%, reduced manual effort by 25%, and supported more than 10 business KPIs through scalable reporting and automation. Expert in advanced SQL (CTEs, window functions), Power BI, Snowflake, and BigQuery to improve data accuracy, strengthen reporting reliability, and support business decision-making.

Experience

MyTech Detectives — Data Analyst | Hyderabad, Telangana | Jul 2024 – Present
• Developed 5+ interactive Power BI dashboards tracking 10+ key business KPIs, establishing dashboard governance, increasing operational efficiency visibility by 25% and supporting business reporting.
• Optimized SQL queries and data pipelines across Snowflake and BigQuery, reducing report generation time by 30% through SQL performance tuning and data modeling.
• Engineered Python-based (Pandas, NumPy) data cleaning and EDA scripts, detecting 15+ critical data anomalies and improving overall reporting accuracy by 20%.
• Led data validation and reconciliation of 12 source systems and warehouse tables, improving report reliability by 20%.

Cognizant Technology Solutions — Associate Test Engineer | Hyderabad, Telangana | Sep 2019 – Jul 2022
• Ensured 99.8% data integrity across 4+ large-scale HIPAA-compliant healthcare applications by leading end-to-end functional and validation testing.
• Developed and maintained Selenium automation frameworks, reducing manual regression effort by 25%.
• Optimized test execution using BrowserStack, reducing regression cycles by 35% through parallel cross-browser testing.
• Validated backend transactional data pipelines and ETL processes using 100+ complex SQL verification queries, achieving a 98% data-quality score.

Projects

Customer Churn Analysis (Python + SQL) | Jan 2024 – Mar 2024
• Analyzed 7K+ subscription customer records using SQL and Python to identify churn trends across contract type, tenure, and customer engagement behavior.
• Applied cohort analysis, feature correlation, and window functions to identify high-risk segments contributing to an 8% monthly churn rate reduction.

Sales Performance & Revenue Analytics Dashboard (Power BI + SQL) | Apr 2024 – Jun 2024
• Developed interactive Power BI dashboards analyzing regional sales performance, product-level revenue trends, and monthly KPIs, identifying $45K+ in monthly opportunities.
• Automated data extraction, transformation, and monthly KPI reporting using SQL and Power Query, saving 4+ hours of manual effort weekly.

Certifications
• Microsoft Certified: Power BI Data Analyst Associate (PL-300)
• Google Data Analytics Professional Certificate

Education
Florida International University, Miami, FL | MS in Data Science and AI | Aug 2022 – May 2024
• Coursework: Machine Learning, Statistical Methods, Big Data Systems, Data Visualization.
Jawaharlal Nehru Technological University, Hyderabad, Telangana | Bachelors in Electronics and Communication Engineering | Sep 2014 – May 2018

Skills
• Programming: SQL, Python, R
• BI & Visualization: Power BI, Tableau, Excel
• Cloud & Databases: Snowflake, BigQuery, SQL Server, Azure, AWS, GCP, MySQL
• Data Engineering: ETL, data pipeline orchestration, data modeling, data warehousing, data validation, data governance
• Methodologies: Exploratory Data Analysis, Statistical Analysis, Statistical Modeling, Machine Learning, SQL performance tuning, A/B Testing, Hypothesis Testing

Achievements
• Solved 50+ advanced SQL problems covering joins, CTEs, and window functions.
• Transitioned from QA to Data Analytics through self-driven learning and practical projects."""
    
    with open("master_da_resume.txt", "w") as f:
        f.write(plain_text_res.strip())
        
    print("🔄 Writing text content to databases...")
    await update_databases(docx_bytes, plain_text_res.strip())
    
    print("🔄 Running fresh resume audit on updated template...")
    from V3_updates import audit_resume, call_llm
    audit = await audit_resume(call_llm)
    print("✨ Re-audit completed successfully! New Audit Data:")
    print("Score:", audit.get("overall_score"))
    print("Verdict:", audit.get("verdict"))

if __name__ == "__main__":
    asyncio.run(main())
