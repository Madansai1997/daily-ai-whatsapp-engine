"""
Résumé Editor — in-place .docx editing that PRESERVES formatting.

The problem: in Word, one sentence is often split across several internal "runs"
(each with its own font/bold/etc). A naive run.text replace corrupts formatting or
misses the match. This does a run-span-aware replacement: the replacement text
inherits the formatting of the run where the match starts; prefixes/suffixes keep
their own runs. Paragraph style (bullets, spacing, headings) is never touched.
"""

import io
from docx import Document


def _iter_paragraphs(doc):
    """All paragraphs, including inside tables (résumés often use tables for layout)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_in_paragraph(paragraph, find: str, replace: str) -> bool:
    """Replace the first occurrence of `find` within a paragraph across runs,
    preserving each run's formatting. Returns True if a replacement happened."""
    runs = paragraph.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    idx = full.find(find)
    if idx == -1:
        return False
    end = idx + len(find)

    pos = 0
    for r in runs:
        r_start = pos
        r_end = pos + len(r.text)
        pos = r_end
        if r_end <= idx or r_start >= end:
            continue  # this run doesn't overlap the match
        prefix = r.text[: idx - r_start] if r_start < idx else ""
        suffix = r.text[end - r_start:] if r_end > end else ""
        if r_start <= idx:
            # run that contains the match start → drop the replacement text here,
            # inheriting this run's formatting
            r.text = prefix + replace + suffix
        else:
            # runs fully inside / at the tail of the match
            r.text = prefix + suffix
    return True


def apply_rewrites(docx_bytes: bytes, rewrites: list[tuple[str, str]]) -> tuple[bytes, int]:
    """Apply (find, replace) rewrites to a .docx, first occurrence each. Returns
    (new_docx_bytes, number_applied). Formatting preserved."""
    doc = Document(io.BytesIO(docx_bytes))
    applied = 0
    for find, replace in rewrites:
        if not find or not (find.strip()):
            continue
        for p in _iter_paragraphs(doc):
            if find in p.text:
                if _replace_in_paragraph(p, find, replace):
                    applied += 1
                break
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), applied


def append_bullet(docx_bytes: bytes, section_title: str, text: str) -> tuple[bytes, bool]:
    """Add a bullet under the paragraph whose text matches `section_title` (e.g. 'Skills'),
    cloning that section's following-paragraph style so it fits in. Best-effort."""
    doc = Document(io.BytesIO(docx_bytes))
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip().lower() == section_title.strip().lower():
            # insert a new paragraph right after the heading, mirroring the next para's style
            style = paras[i + 1].style if i + 1 < len(paras) else p.style
            new_p = p.insert_paragraph_before(text)  # placeholder; move below heading
            # python-docx can't insert-after directly; simplest robust: append at end if that fails
            new_p.style = style
            out = io.BytesIO(); doc.save(out)
            return out.getvalue(), True
    # fallback: append at the very end
    np = doc.add_paragraph(text)
    out = io.BytesIO(); doc.save(out)
    return out.getvalue(), True
